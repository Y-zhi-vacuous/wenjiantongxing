"""生成创AI比赛报告 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# 默认样式
style = doc.styles['Normal']
style.font.name = '仿宋_GB2312'
style.font.size = Pt(16)
style.paragraph_format.line_spacing = Pt(28)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

with open('docs/开发与应用报告.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

def set_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

for line in lines:
    text = line.rstrip()
    if not text:
        continue

    # 主标题
    if text.startswith('# ') and '报告' in text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text[2:]), '方正小标宋简体', 18)
        continue

    # 一级标题
    if text.startswith('## '):
        p = doc.add_paragraph()
        set_font(p.add_run(text[3:]), '黑体', 16, bold=True)
        continue

    # 二级标题
    if text.startswith('### '):
        p = doc.add_paragraph()
        set_font(p.add_run(text[4:]), '楷体_GB2312', 16)
        continue

    # 三级标题
    if text.startswith('#### '):
        p = doc.add_paragraph()
        set_font(p.add_run(text[5:]), '仿宋_GB2312', 16, bold=True)
        continue

    # 跳过代码块
    if text.startswith('```'):
        continue

    # 分隔线
    if text == '---':
        continue

    # 表格行
    if text.startswith('|') and '---' not in text:
        cells = [c.strip() for c in text.split('|')[1:-1]]
        try:
            table = doc.add_table(rows=1, cols=len(cells), style='Table Grid')
            for i, ct in enumerate(cells):
                table.rows[0].cells[i].text = ct
        except:
            pass
        continue

    # 带 `代码` 的行
    if '`' in text:
        p = doc.add_paragraph()
        parts = re.split(r'(`[^`]+`)', text)
        for part in parts:
            if part.startswith('`') and part.endswith('`'):
                set_font(p.add_run(part[1:-1]), '仿宋_GB2312', 14)
            else:
                set_font(p.add_run(part), '仿宋_GB2312', 16)
        continue

    # 粗体标记
    if '**' in text:
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                set_font(p.add_run(part[2:-2]), '仿宋_GB2312', 16, bold=True)
            else:
                set_font(p.add_run(part), '仿宋_GB2312', 16)
        continue

    # 引用
    if text.startswith('> '):
        p = doc.add_paragraph()
        set_font(p.add_run(text[2:]), '仿宋_GB2312', 16)
        continue

    # 普通正文
    if text.startswith('- '):
        p = doc.add_paragraph()
        set_font(p.add_run(text), '仿宋_GB2312', 16)
    else:
        p = doc.add_paragraph()
        set_font(p.add_run(text), '仿宋_GB2312', 16)

doc.save('docs/文鉴同行-创AI案例开发与应用报告.docx')
print('Done: docs/文鉴同行-创AI案例开发与应用报告.docx')
